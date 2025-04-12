from flask import Flask, render_template, request, jsonify, redirect, url_for
import os
from werkzeug.utils import secure_filename
from pdfminer.high_level import extract_text
import docx
import re
import io

app = Flask(__name__)
UPLOAD_FOLDER = "uploads"
ALLOWED_EXTENSIONS = {"pdf", "docx"}
MAX_RESUMES = 10

app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

# Ensure uploads folder exists
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# Temporary storage for file data
selected_files = []

# Check allowed file types
def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

# Extract text from PDF
def extract_text_from_pdf(file_path):
    return extract_text(file_path)

# Extract text from DOCX
def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

# Extract CGPA/percentage and calculate average
def extract_and_average_cgpa(text):
    cgpa_matches = re.findall(r"\b(?:CGPA|cgpa|Cgpa)\s*[:=]?\s*(\d+(?:\.\d+)?)", text)
    percentage_matches = re.findall(r"\b(\d{1,3})\s*%", text)

    values = []

    # Convert matches to float
    values.extend([float(match) for match in cgpa_matches])
    values.extend([float(match) for match in percentage_matches if float(match) <= 100])

    # Return average if values exist, otherwise 0
    return sum(values) / len(values) if values else 0

@app.route("/", methods=["GET", "POST"])
def index():
    global selected_files
    candidates = []

    if request.method == "POST":
        # Add file to temporary list
        if "resume" in request.files:
            file = request.files["resume"]

            if file and allowed_file(file.filename):
                # Store file data as bytes to avoid closure issues
                file_data = {
                    "filename": secure_filename(file.filename),
                    "content": file.read()
                }
                selected_files.append(file_data)

        # Upload all selected files
        if request.form.get("upload_all"):
            if not selected_files:
                return "No resumes selected for upload"

            if len(selected_files) > MAX_RESUMES:
                return f"Maximum {MAX_RESUMES} resumes allowed"

            # Process selected files
            for file_data in selected_files:
                filename = file_data["filename"]
                file_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)

                # Save file from bytes
                with open(file_path, "wb") as f:
                    f.write(file_data["content"])

                # Extract text
                if filename.endswith(".pdf"):
                    extracted_text = extract_text_from_pdf(file_path)
                elif filename.endswith(".docx"):
                    with io.BytesIO(file_data["content"]) as docx_file:
                        doc = docx.Document(docx_file)
                        extracted_text = "\n".join([para.text for para in doc.paragraphs])
                else:
                    continue
                # Calculate CGPA/percentage
                avg_cgpa = extract_and_average_cgpa(extracted_text)

                candidates.append({
                    "filename": filename,
                    "cgpa": avg_cgpa,
                    "text": extracted_text
                })

            # Sort and rank candidates
            candidates.sort(key=lambda x: x["cgpa"], reverse=True)
            for rank, candidate in enumerate(candidates, start=1):
                candidate["rank"] = rank

            # Clear temporary list after processing
            selected_files.clear()

    return render_template("index.html", candidates=candidates, selected_files=selected_files)

if __name__ == "__main__":
    app.run(debug=True)
